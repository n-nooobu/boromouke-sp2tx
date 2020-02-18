import os
import pydub
from docx import Document

# Imports the Google Cloud client library
from google.cloud import storage
from google.cloud import speech
from google.cloud.speech import enums
from google.cloud.speech import types

file_name = 'trial'

# flacに変換
sound = pydub.AudioSegment.from_file('sound/' + file_name + '.mp3', 'mp3')
sound = sound.set_channels(1)
sound = sound.set_frame_rate(16000)
sound.export('sound/' + file_name + '.flac', format='flac')

# 環境変数設定 私のvast-box-268016-0698f0b37e52というGCPプロジェクトを使いますよという宣言
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = './vast-box-268016-0698f0b37e52.json'

# Upload to Google Cloud Storage(1分以上の音声をテキストに変換する場合はGCSに配置しなきゃダメ)
client_sto = storage.Client()
bucket = client_sto.get_bucket('boromouke-sp2tx')
blob = bucket.blob(file_name + '.flac')
blob.upload_from_filename('sound/' + file_name + '.flac')

client = speech.SpeechClient()

audio = types.RecognitionAudio(uri='gs://boromouke-sp2tx/' + file_name + '.flac')
config = types.RecognitionConfig(
    encoding=enums.RecognitionConfig.AudioEncoding.FLAC,
    sample_rate_hertz=16000,
    language_code='ja-JP')

operation = client.long_running_recognize(config, audio)

print('Waiting for operation to complete...')
response = operation.result()

# テンプレワードファイルを読み込みます
document = Document('word/template.docx')

for result in response.results:
    document.paragraphs[0].add_run(result.alternatives[0].transcript)

document.save('word/' + file_name + '.docx')